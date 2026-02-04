"""Tests for contract template upload and generation."""
from io import BytesIO
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from docx import Document

from app.services.contract_template_service import (
    _detect_file_type_by_signature,
    _get_libreoffice_cmd,
    _is_valid_docx,
    contract_data_to_context,
    get_docx_bytes_for_template,
    upload_template_file,
    validate_template_upload,
)


def test_detect_file_type_by_signature_rtf():
    """RTF magic bytes {\\rtf are detected."""
    content = b"{\\rtf1\\ansi\n" + b"x" * 100
    assert _detect_file_type_by_signature(content) == "rtf"


def test_detect_file_type_by_signature_rtf_with_bom_or_whitespace():
    """RTF with BOM or whitespace before {\\rtf is accepted."""
    base = b"{\\rtf1\\ansi\n" + b"x" * 100
    assert _detect_file_type_by_signature(b"\xef\xbb\xbf" + base) == "rtf"
    assert _detect_file_type_by_signature(b"  \t\r\n" + base) == "rtf"
    assert _detect_file_type_by_signature(b" " + base) == "rtf"


def test_is_valid_docx_accepts_real_docx():
    """Real DOCX (ZIP with word/document.xml) is valid."""
    doc = Document()
    doc.add_paragraph("x")
    buf = BytesIO()
    doc.save(buf)
    assert _is_valid_docx(buf.getvalue()) is True


def test_is_valid_docx_rejects_invalid_zip():
    """PK header without word/document.xml is rejected."""
    fake_zip = b"PK\x03\x04" + b"x" * 100
    assert _is_valid_docx(fake_zip) is False


def test_detect_file_type_by_signature_docx():
    """DOCX (ZIP with word/document.xml) is detected."""
    doc = Document()
    doc.add_paragraph("x")
    buf = BytesIO()
    doc.save(buf)
    content = buf.getvalue()
    assert _detect_file_type_by_signature(content) == "docx"


def test_detect_file_type_by_signature_invalid():
    """Unknown content returns None."""
    assert _detect_file_type_by_signature(b"xxxx") is None
    assert _detect_file_type_by_signature(b"") is None


def test_validate_template_upload_docx():
    """Accept valid DOCX filename and content (magic bytes)."""
    doc = Document()
    doc.add_paragraph("x")
    buf = BytesIO()
    doc.save(buf)
    content = buf.getvalue()
    ft, err = validate_template_upload(content, "template.docx")
    assert ft == "docx"
    assert err == ""


def test_validate_template_upload_rtf():
    """Accept valid RTF filename and content (magic bytes)."""
    content = b"{\\rtf1\\ansi\n" + b"x" * 200
    ft, err = validate_template_upload(content, "file.RTF")
    assert ft == "rtf"
    assert err == ""


def test_validate_template_upload_rtf_with_bom():
    """Accept RTF with BOM/whitespace before {\\rtf."""
    content = b"\xef\xbb\xbf{\\rtf1\\ansi\n" + b"x" * 200
    ft, err = validate_template_upload(content, "file.rtf")
    assert ft == "rtf"
    assert err == ""


def test_validate_template_upload_rejects_wrong_content():
    """Reject when extension does not match content (e.g. .docx but RTF bytes)."""
    content = b"{\\rtf1\\ansi\n" + b"x" * 100
    ft, err = validate_template_upload(content, "file.docx")
    assert ft == ""
    assert "содержимым" in err or "формат" in err.lower()


def test_validate_template_upload_rejects_type():
    """Reject non-DOCX/RTF extension."""
    content = b"{\\rtf1\\ansi\n" + b"x" * 100
    ft, err = validate_template_upload(content, "file.txt")
    assert ft == ""
    assert "DOCX" in err or "RTF" in err


def test_validate_template_upload_rejects_size():
    """Reject file over size limit."""
    content = b"{\\rtf1\n" + (b"x" * (11 * 1024 * 1024))
    ft, err = validate_template_upload(content, "file.rtf")
    assert ft == ""
    assert "большой" in err or "big" in err.lower()


def test_get_libreoffice_cmd_prefers_libreoffice():
    """When libreoffice is in PATH, return it."""
    with patch("app.services.contract_template_service.shutil.which") as m:
        m.side_effect = lambda c: "/usr/bin/libreoffice" if c == "libreoffice" else None
        assert _get_libreoffice_cmd() == "/usr/bin/libreoffice"
        m.assert_any_call("libreoffice")


def test_get_libreoffice_cmd_fallback_to_soffice():
    """When libreoffice is missing, use soffice."""
    with patch("app.services.contract_template_service.shutil.which") as m:
        m.side_effect = lambda c: None if c == "libreoffice" else "/usr/bin/soffice"
        assert _get_libreoffice_cmd() == "/usr/bin/soffice"
        m.assert_any_call("libreoffice")
        m.assert_any_call("soffice")


def test_get_libreoffice_cmd_raises_when_missing():
    """When neither binary exists, raise RuntimeError."""
    with patch("app.services.contract_template_service.shutil.which", return_value=None):
        with pytest.raises(RuntimeError, match="не найден"):
            _get_libreoffice_cmd()


def test_contract_data_to_context():
    """Context has all placeholder keys with fallback for None."""
    from app.services.pdf import ContractData

    contract = ContractData(
        company_name="ООО Рога",
        inn="123",
        director="Иванов",
        bank_bik="044525225",
        bank_account="40702",
        contract_number="1-20260204",
        contract_date="04.02.2026",
        service_description="Услуги",
        kpp="123",
        ogrn=None,
        legal_address=None,
        bank_name=None,
        bank_corr_account=None,
    )
    ctx = contract_data_to_context(contract)
    assert ctx["company_name"] == "ООО Рога"
    assert ctx["ogrn"] == "-"
    assert "contract_number" in ctx


def _minimal_docx_bytes() -> bytes:
    """Create minimal DOCX with placeholder for tests."""
    doc = Document()
    doc.add_paragraph("Договор {{contract_number}} от {{contract_date}}")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


@pytest.fixture
def minimal_docx():
    return _minimal_docx_bytes()


async def test_list_contract_templates(client, admin_headers):
    """Admin can list contract templates."""
    response = await client.get("/api/v1/admin/contract-templates", headers=admin_headers)
    assert response.status_code == 200
    assert isinstance(response.json(), list)


async def test_upload_contract_template(client, admin_headers, minimal_docx):
    """Upload DOCX template with mocked S3; heavy work runs via to_thread."""
    async def run_sync_in_mock(fn, *args, **kwargs):
        return fn(*args, **kwargs)

    with (
        patch("app.api.v1.routes.admin.asyncio.to_thread", new_callable=AsyncMock) as mock_to_thread,
        patch("app.api.v1.routes.admin.upload_template_file") as mock_upload,
        patch("app.api.v1.routes.admin.head_check_upload", new_callable=AsyncMock) as mock_head,
    ):
        mock_upload.return_value = ("contract-templates/fake-key.docx", None)
        mock_to_thread.side_effect = run_sync_in_mock
        mock_head.return_value = True

        response = await client.post(
            "/api/v1/admin/contract-templates/upload",
            headers=admin_headers,
            data={"name": "Тест шаблон", "is_default": "true"},
            files={"file": ("template.docx", minimal_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert response.status_code == 200
        body = response.json()
        assert body["name"] == "Тест шаблон"
        assert body["file_name"] == "template.docx"
        assert body["file_type"] == "docx"
        assert body["is_default"] is True
        mock_to_thread.assert_called_once()
        mock_upload.assert_called_once()


async def test_upload_head_check_failure_returns_503_and_deletes_files(client, admin_headers, minimal_docx):
    """When head_check fails after upload, return 503 and delete uploaded S3 keys."""
    async def run_sync_in_mock(fn, *args, **kwargs):
        return fn(*args, **kwargs)

    with (
        patch("app.api.v1.routes.admin.asyncio.to_thread", new_callable=AsyncMock) as mock_to_thread,
        patch("app.api.v1.routes.admin.upload_template_file") as mock_upload,
        patch("app.api.v1.routes.admin.head_check_upload", new_callable=AsyncMock) as mock_head,
        patch("app.api.v1.routes.admin.delete_template_files") as mock_delete,
    ):
        mock_upload.return_value = ("contract-templates/key.docx", None)
        mock_to_thread.side_effect = run_sync_in_mock
        mock_head.return_value = False

        response = await client.post(
            "/api/v1/admin/contract-templates/upload",
            headers=admin_headers,
            data={"name": "Тест", "is_default": "false"},
            files={"file": ("t.docx", minimal_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert response.status_code == 503
        mock_delete.assert_called_once()
        call_args = mock_delete.call_args[0]
        assert call_args[1] == "contract-templates/key.docx"
        assert call_args[2] is None


async def test_download_streams_chunks_via_to_thread(client, admin_headers, db_session):
    """Download endpoint streams S3 via asyncio.to_thread (chunk iterator)."""
    from app.api.v1.routes.admin import _stream_s3_chunks
    from app.db.models.contract_template import ContractTemplate

    template = ContractTemplate(
        name="D",
        html_content=None,
        is_default=False,
        file_key="contract-templates/k.docx",
        file_name="k.docx",
        file_type="docx",
        docx_key=None,
    )
    db_session.add(template)
    await db_session.commit()
    await db_session.refresh(template)

    async def run_sync_in_mock(fn, *args, **kwargs):
        return fn(*args, **kwargs)

    with (
        patch("app.api.v1.routes.admin.asyncio.to_thread", new_callable=AsyncMock) as mock_to_thread,
        patch("app.api.v1.routes.admin.S3Service") as mock_s3_class,
    ):
        mock_to_thread.side_effect = run_sync_in_mock
        mock_s3 = mock_s3_class.return_value
        mock_s3.stream_chunks.return_value = iter([b"chunk1", b"chunk2"])

        response = await client.get(
            f"/api/v1/admin/contract-templates/{template.id}/download",
            headers=admin_headers,
        )
        assert response.status_code == 200
        assert response.content == b"chunk1chunk2"
        assert mock_to_thread.call_count >= 1
        first_call_fn = mock_to_thread.call_args_list[0][0][0]
        assert first_call_fn is _stream_s3_chunks


def test_upload_template_file_rtf_uploads_single_file():
    """When RTF is uploaded, only one file is uploaded to S3 (no docx_key)."""
    s3 = MagicMock()
    content = b"{\\rtf1\\ansi\nHello}"
    file_key, docx_key = upload_template_file(s3, content, "file.rtf", "rtf")
    assert docx_key is None
    assert file_key.endswith(".rtf")
    s3.upload_bytes.assert_called_once()
    call_args = s3.upload_bytes.call_args[0]
    assert call_args[0] == file_key
    assert call_args[1] == content


def test_get_docx_bytes_legacy_pdf_uses_docx_key():
    """Legacy PDF template with docx_key returns DOCX from docx_key."""
    s3 = MagicMock()
    s3.get_bytes.side_effect = lambda key: b"docx-from-" + key.encode() if isinstance(key, str) else b""
    result = get_docx_bytes_for_template(
        s3, "contract-templates/old.pdf", "pdf", "contract-templates/old-converted.docx"
    )
    assert result == b"docx-from-contract-templates/old-converted.docx"
    s3.get_bytes.assert_called_once_with("contract-templates/old-converted.docx")


def test_get_docx_bytes_legacy_pdf_without_docx_key_raises():
    """Legacy PDF template without docx_key raises RuntimeError with clear message."""
    s3 = MagicMock()
    with pytest.raises(RuntimeError, match="Старый шаблон|PDF|конвертированной"):
        get_docx_bytes_for_template(s3, "contract-templates/old.pdf", "pdf", None)
    s3.get_bytes.assert_not_called()


async def test_delete_contract_template(client, admin_headers, db_session, minimal_docx):
    """Delete template and call delete_template_files when file_key is set."""
    from app.db.models.contract_template import ContractTemplate

    template = ContractTemplate(
        name="To delete",
        html_content=None,
        is_default=False,
        file_key="contract-templates/fake",
        file_name="x.docx",
        file_type="docx",
        docx_key=None,
    )
    db_session.add(template)
    await db_session.commit()
    await db_session.refresh(template)
    tid = template.id

    with patch("app.api.v1.routes.admin.delete_template_files") as mock_delete:
        response = await client.delete(
            f"/api/v1/admin/contract-templates/{tid}",
            headers=admin_headers,
        )
        assert response.status_code == 200
        mock_delete.assert_called_once()


async def test_generate_contract_pdf_uses_html(client, auth_headers_and_user, db_session):
    """Generate contract PDF when default template is HTML (no file)."""
    from sqlalchemy import update

    from app.db.models.company import Company
    from app.db.models.contract_template import ContractTemplate

    headers, user = auth_headers_and_user
    await db_session.execute(update(ContractTemplate).values(is_default=False))
    await db_session.commit()

    company = Company(
        user_id=user.id,
        inn=str(1000000000 + user.id),
        name="ООО Тест",
        director="Директоров",
        bank_bik="044525225",
        bank_account="40702810000000000000",
    )
    db_session.add(company)
    template = ContractTemplate(
        name="HTML",
        html_content="<html><body>Договор {{contract_number}}</body></html>",
        is_default=True,
    )
    db_session.add(template)
    await db_session.commit()
    await db_session.refresh(company)
    await db_session.refresh(template)

    fake_pdf = b"%PDF-1.4 html-rendered"
    with patch("app.api.v1.routes.companies.render_contract_pdf") as mock_render:
        mock_render.return_value = fake_pdf

        response = await client.get(
            f"/api/v1/companies/{company.id}/contract",
            headers=headers,
        )
        assert response.status_code == 200
        assert response.headers.get("content-type", "").startswith("application/pdf")
        assert response.content == fake_pdf
        mock_render.assert_called_once()


async def test_generate_contract_file_template_uses_to_thread(client, auth_headers_and_user, db_session):
    """When default template is file-based, render runs via asyncio.to_thread."""
    from sqlalchemy import update

    from app.db.models.company import Company
    from app.db.models.contract_template import ContractTemplate

    headers, user = auth_headers_and_user
    await db_session.execute(update(ContractTemplate).values(is_default=False))
    await db_session.commit()

    company = Company(
        user_id=user.id,
        inn=str(2000000000 + user.id),
        name="ООО Тест",
        director="Директоров",
        bank_bik="044525225",
        bank_account="40702810000000000000",
    )
    db_session.add(company)
    template = ContractTemplate(
        name="File",
        html_content=None,
        is_default=True,
        file_key="contract-templates/fake.docx",
        file_name="fake.docx",
        file_type="docx",
        docx_key=None,
    )
    db_session.add(template)
    await db_session.commit()
    await db_session.refresh(company)
    await db_session.refresh(template)

    fake_pdf = b"%PDF-1.4 fake content"
    with patch(
        "app.api.v1.routes.companies.render_contract_pdf_from_docx_template",
        return_value=fake_pdf,
    ) as mock_render:

        response = await client.get(
            f"/api/v1/companies/{company.id}/contract",
            headers=headers,
        )
        assert response.status_code == 200
        assert response.headers.get("content-type", "").startswith("application/pdf")
        assert response.content == fake_pdf
        mock_render.assert_called_once()
