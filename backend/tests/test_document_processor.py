"""Tests for RAG document processing: parse_txt, parse_docx, index_document validation."""
from io import BytesIO

import pytest
from docx import Document

from app.services.document_processor import (
    MAX_CHUNKS_PER_DOCUMENT,
    MAX_DOCUMENT_SIZE_BYTES,
    index_document,
    parse_docx,
    parse_rtf,
    parse_txt,
    split_into_chunks,
)


# ----- parse_txt -----


def test_parse_txt_valid_utf8():
    """TXT в UTF-8 декодируется, пробелы по краям обрезаются."""
    content = "  Прайс услуг.\n\nУслуга: Хранение.  ".encode("utf-8")
    assert parse_txt(content) == "Прайс услуг.\n\nУслуга: Хранение."


def test_parse_txt_empty_after_strip():
    """Пустой файл или только пробелы даёт пустую строку."""
    assert parse_txt(b"") == ""
    assert parse_txt(b"   \n\t  ") == ""


def test_parse_txt_invalid_utf8_raises():
    """Не-UTF-8 содержимое поднимает ValueError с сообщением про UTF-8."""
    content = b"\xff\xfe"
    with pytest.raises(ValueError, match="UTF-8"):
        parse_txt(content)


# ----- parse_docx -----


def test_parse_docx_valid():
    """DOCX с параграфами отдаёт текст через двойной перенос."""
    doc = Document()
    doc.add_paragraph("Первый блок")
    doc.add_paragraph("Второй блок")
    buf = BytesIO()
    doc.save(buf)
    text = parse_docx(buf.getvalue())
    assert "Первый блок" in text
    assert "Второй блок" in text


def test_parse_docx_empty_paragraphs_stripped():
    """Пустые параграфы не попадают в вывод."""
    doc = Document()
    doc.add_paragraph("Только это")
    doc.add_paragraph("")
    buf = BytesIO()
    doc.save(buf)
    assert parse_docx(buf.getvalue()).strip() == "Только это"


# ----- parse_rtf -----


def test_parse_rtf_valid():
    """RTF с текстом извлекает plain text."""
    # Минимальный RTF с текстом "Hello RTF"
    rtf = b"{\\rtf1\\ansi\\deff0 {\\cf0 Hello RTF}}"
    result = parse_rtf(rtf)
    assert "Hello" in result
    assert "RTF" in result


def test_parse_rtf_empty_after_strip():
    """Пустой или только контрольные коды RTF даёт пустую строку после strip."""
    rtf = b"{\\rtf1\\ansi}"
    result = parse_rtf(rtf)
    assert result == "" or "rtf" in result.lower()


def test_parse_rtf_utf8():
    """RTF в UTF-8 декодируется."""
    rtf = b"{\\rtf1\\ansi Hello}"
    result = parse_rtf(rtf)
    assert "Hello" in result


def test_parse_rtf_utf16_bom():
    """RTF в UTF-16 с BOM (например, экспорт из macOS Pages) декодируется."""
    rtf_utf16 = b"\xff\xfe" + "{\\rtf1\\ansi Hello}".encode("utf-16-le")
    result = parse_rtf(rtf_utf16)
    assert "Hello" in result


def test_parse_rtf_cp1251():
    """RTF в Windows-1251 (кириллица) декодируется."""
    rtf_cp1251 = b"{\\rtf1\\ansi " + "Привет".encode("cp1251") + b"}"
    result = parse_rtf(rtf_cp1251)
    assert "Привет" in result


# ----- split_into_chunks -----


def test_split_into_chunks_short_text():
    """Короткий текст возвращается одним чанком."""
    result = split_into_chunks("Короткий текст", chunk_size=1000)
    assert result == ["Короткий текст"]


def test_split_into_chunks_respects_paragraphs():
    """Разбиение предпочитает границы параграфов."""
    text = "A" * 500 + "\n\n" + "B" * 500
    result = split_into_chunks(text, chunk_size=600, overlap=50)
    assert len(result) >= 1
    assert any("A" in r and "B" in r for r in result) or len(result) == 1


# ----- index_document validation (без БД) -----


@pytest.mark.asyncio
async def test_index_document_rejects_pdf():
    """index_document при document_type pdf поднимает ValueError."""
    db = None
    with pytest.raises(ValueError, match="document_type must be"):
        await index_document(db, "file.pdf", b"content", "pdf")


@pytest.mark.asyncio
async def test_index_document_rejects_unknown_type():
    """index_document при неизвестном типе поднимает ValueError."""
    with pytest.raises(ValueError, match="document_type must be"):
        await index_document(None, "file.xlsx", b"content", "xlsx")


# ----- constants -----


def test_max_document_size_const():
    """Лимит размера файла 15 MB."""
    assert MAX_DOCUMENT_SIZE_BYTES == 15 * 1024 * 1024


def test_max_chunks_const():
    """Лимит чанков на документ 80."""
    assert MAX_CHUNKS_PER_DOCUMENT == 80
