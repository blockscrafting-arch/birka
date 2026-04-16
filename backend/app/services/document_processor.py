"""Document processing for RAG: DOCX/TXT parsing, chunking, indexing."""
import json
import re
from datetime import datetime, timezone
from io import BytesIO

from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.sql import func

from app.core.config import settings
from app.core.logging import logger
from app.db.models.document_chunk import DocumentChunk
from app.services.rag import get_embedding

# Ограничения, чтобы не ронять воркер при больших файлах и множестве эмбеддингов
MAX_DOCUMENT_SIZE_BYTES = 15 * 1024 * 1024  # 15 MB
MAX_CHUNKS_PER_DOCUMENT = 80

# RTF \ansicpg code page -> Python encoding name (common Windows/Office code pages)
_ANSI_CPG_TO_ENCODING: dict[int, str] = {
    1250: "cp1250",  # Central Europe
    1251: "cp1251",  # Cyrillic
    1252: "cp1252",  # Western Europe
    1254: "cp1254",  # Turkish
    1257: "cp1257",  # Baltic
}


def _get_rtf_ansicpg(content: bytes) -> int | None:
    """Extract \\ansicpgN from RTF header (first 2 KB). Returns code page number or None."""
    header = content[:2048]
    match = re.search(rb"\\ansicpg(\d+)", header)
    if match:
        return int(match.group(1))
    return None


def parse_txt(content: bytes) -> str:
    """Decode plain text (UTF-8)."""
    try:
        return content.decode("utf-8").strip()
    except UnicodeDecodeError as exc:
        logger.warning("txt_decode_failed", error=str(exc))
        raise ValueError("Файл должен быть в кодировке UTF-8") from exc


def parse_docx(content: bytes) -> str:
    """Extract text from DOCX. Returns concatenated paragraph text."""
    try:
        from docx import Document

        doc = Document(BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(parts).strip()
    except Exception as exc:
        logger.exception("docx_parse_failed", error=str(exc))
        raise


# Section title (from DOCX) -> section key for RAG filter (fallback if config empty)
DOCX_SECTION_HEADING_TO_KEY: dict[str, str] = {
    "Обувь": "obuv",
    "Зеркала": "zerkala",
    "Посуда": "posuda",
    "Упаковка хрупких товаров": "hrupkie",
    "Сантехника": "santehnika",
    "Автомобильные масла": "avto_masla",
    "Автомобильные аккумуляторы": "avto_akkum",
    "Мебель": "mebel",
    "Матрасы и топперы": "matrasy",
    "Одежда и аксессуары": "odezhda",
    "Одежда, обувь, аксессуары": "odezhda",
}


def get_docx_section_heading_to_key() -> dict[str, str]:
    """Маппинг заголовок -> section. Из конфига RAG_DOCX_SECTION_HEADINGS_JSON или дефолт из кода."""
    raw = getattr(settings, "RAG_DOCX_SECTION_HEADINGS_JSON", "") or ""
    if not raw.strip():
        return DOCX_SECTION_HEADING_TO_KEY.copy()
    try:
        data = json.loads(raw)
        if isinstance(data, dict) and all(isinstance(k, str) and isinstance(v, str) for k, v in data.items()):
            return data
    except (json.JSONDecodeError, TypeError):
        pass
    return DOCX_SECTION_HEADING_TO_KEY.copy()


def _normalize_heading(text: str) -> str:
    return (text or "").strip()


def parse_docx_with_sections(content: bytes) -> list[tuple[str, str | None]]:
    """
    Parse DOCX into chunks with section keys. Uses paragraph styles (Heading 1/2) or
    fallback to known titles. Each chunk is (text, section_key). section_key None for intro/unknown.
    """
    try:
        from docx import Document

        doc = Document(BytesIO(content))
    except Exception as exc:
        logger.exception("docx_parse_failed", error=str(exc))
        raise

    mapping = get_docx_section_heading_to_key()
    result: list[tuple[str, str | None]] = []
    current_section: str | None = None
    current_text: list[str] = []
    chunk_size = 900
    overlap = 150

    def flush_section() -> None:
        nonlocal current_text
        if not current_text:
            return
        full = "\n\n".join(current_text).strip()
        if not full:
            return
        sub_chunks = split_into_chunks(full, chunk_size=chunk_size, overlap=overlap)
        for sub in sub_chunks:
            if sub:
                result.append((sub, current_section))
        current_text = []

    for p in doc.paragraphs:
        text = _normalize_heading(p.text)
        style_name = getattr(p.style, "name", None) or ""
        is_heading = isinstance(style_name, str) and style_name.startswith("Heading")

        # Heading by style or by exact match with known section title (fallback when no Heading style)
        known_section = mapping.get(text)
        if known_section is None and text and is_heading:
            for title, key in mapping.items():
                if title in text or text in title:
                    known_section = key
                    break
        is_section_header = (is_heading and text) or (bool(text) and text in mapping)

        if is_section_header:
            flush_section()
            current_section = mapping.get(text) if text in mapping else known_section
            current_text = [text]
        elif text:
            current_text.append(text)

    flush_section()

    if not result:
        fallback_text = parse_docx(content)
        if fallback_text:
            for chunk in split_into_chunks(fallback_text, chunk_size=1000, overlap=200):
                if chunk:
                    result.append((chunk, None))
    return result


def parse_rtf(content: bytes) -> str:
    """
    Extract plain text from RTF. Uses \\ansicpg when present, else UTF-8/UTF-16 (BOM)/cp1251/cp1252.
    """
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError:
        raise ValueError("Обработка RTF недоступна (установите striprtf)") from None
    rtf_str: str
    # 1) UTF-16 BOM
    if content.startswith((b"\xff\xfe", b"\xfe\xff")):
        try:
            rtf_str = content.decode("utf-16")
        except Exception as exc:
            logger.warning("rtf_utf16_decode_failed", error=str(exc))
            rtf_str = ""
        if rtf_str:
            try:
                text = rtf_to_text(rtf_str)
                return (text or "").strip()
            except Exception as exc:
                logger.exception("rtf_parse_failed", error=str(exc))
                raise ValueError("Не удалось извлечь текст из RTF") from exc
    # 2) \ansicpg in header
    cpg = _get_rtf_ansicpg(content)
    if cpg is not None and cpg in _ANSI_CPG_TO_ENCODING:
        enc = _ANSI_CPG_TO_ENCODING[cpg]
        try:
            rtf_str = content.decode(enc)
        except Exception as exc:
            logger.warning("rtf_ansicpg_decode_failed", ansicpg=cpg, encoding=enc, error=str(exc))
            rtf_str = ""
        if rtf_str:
            try:
                text = rtf_to_text(rtf_str)
                return (text or "").strip()
            except Exception as exc:
                logger.exception("rtf_parse_failed", error=str(exc))
                raise ValueError("Не удалось извлечь текст из RTF") from exc
    # 3) Fallback: UTF-8, then cp1251, then cp1252
    try:
        rtf_str = content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            rtf_str = content.decode("cp1251")
        except Exception:
            try:
                rtf_str = content.decode("cp1252")
            except Exception as exc:
                logger.warning("rtf_decode_failed", error=str(exc))
                raise ValueError(
                    "Файл RTF должен быть в кодировке UTF-8, UTF-16 или Windows (ansicpg 1250/1251/1252/1254/1257)"
                ) from exc
    try:
        text = rtf_to_text(rtf_str)
        return (text or "").strip()
    except Exception as exc:
        logger.exception("rtf_parse_failed", error=str(exc))
        raise ValueError("Не удалось извлечь текст из RTF") from exc


def split_into_chunks(
    text: str,
    chunk_size: int = 1000,
    overlap: int = 200,
) -> list[str]:
    """Split text into chunks with overlap. Tries to break at paragraph boundaries."""
    text = (text or "").strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        if end < len(text):
            # Prefer break at paragraph or line end
            search = text.rfind("\n\n", start, end + 1)
            if search > start:
                end = search + 2
            else:
                search = text.rfind("\n", start, end + 1)
                if search > start:
                    end = search + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - overlap
        if start >= len(text):
            break
    return chunks


def _section_for_source_file(source_file: str) -> str | None:
    """Infer section for txt/rtf by filename (wb -> wb_common, ozon -> ozon_common)."""
    lower = (source_file or "").lower()
    if "wb" in lower or "wildberries" in lower:
        return "wb_common"
    if "ozon" in lower:
        return "ozon_common"
    return None


async def index_document(
    db: AsyncSession,
    file_name: str,
    content: bytes,
    document_type: str,
) -> int:
    """
    Parse document, chunk, generate embeddings, store in DB.
    Replaces existing chunks for the same source_file (versioning: next version).
    For DOCX uses section-aware parsing; for txt/rtf assigns wb_common/ozon_common by filename.
    Returns number of chunks added.
    """
    source_file = (file_name or "document").strip() or "document"
    if document_type not in ("docx", "txt", "rtf"):
        raise ValueError("document_type must be 'docx', 'txt' or 'rtf'")
    if len(content) > MAX_DOCUMENT_SIZE_BYTES:
        raise ValueError(
            f"Файл слишком большой (макс. {MAX_DOCUMENT_SIZE_BYTES // (1024*1024)} MB)"
        )

    chunks_with_section: list[tuple[str, str | None]]
    if document_type == "docx":
        chunks_with_section = parse_docx_with_sections(content)
    else:
        if document_type == "txt":
            text = parse_txt(content)
        else:
            text = parse_rtf(content)
        if not text:
            return 0
        chunks_text = split_into_chunks(text, chunk_size=1000, overlap=200)
        section = _section_for_source_file(source_file)
        chunks_with_section = [(c, section) for c in chunks_text if c]

    if not chunks_with_section:
        return 0
    if len(chunks_with_section) > MAX_CHUNKS_PER_DOCUMENT:
        raise ValueError(
            f"Слишком много фрагментов ({len(chunks_with_section)}). "
            f"Максимум {MAX_CHUNKS_PER_DOCUMENT}. Уменьшите размер документа."
        )

    version_result = await db.execute(
        select(func.coalesce(func.max(DocumentChunk.version), 0)).where(
            DocumentChunk.source_file == source_file
        )
    )
    next_version = int(version_result.scalar_one() or 0) + 1

    await db.execute(delete(DocumentChunk).where(DocumentChunk.source_file == source_file))

    now = datetime.now(timezone.utc)
    count = 0
    for i, (chunk_content, section) in enumerate(chunks_with_section):
        embedding = await get_embedding(chunk_content)
        if not embedding:
            continue
        chunk = DocumentChunk(
            content=chunk_content,
            source_file=source_file,
            chunk_index=i,
            embedding=embedding,
            created_at=now,
            document_type=document_type,
            version=next_version,
            section=section,
        )
        db.add(chunk)
        count += 1

    await db.commit()
    logger.info(
        "document_indexed",
        source_file=source_file,
        document_type=document_type,
        version=next_version,
        chunks_added=count,
    )
    return count
